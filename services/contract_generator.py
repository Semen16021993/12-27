import os
from docx import Document
import subprocess
from datetime import datetime


def parse_client_data(case):

    path = f"cases/{case}/passport/client_data.txt"

    data = {}

    with open(path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    for line in lines:

        if ":" in line:

            key, value = line.split(":", 1)

            data[key.strip()] = value.strip()

    return data


def make_signature(fio):

    parts = fio.split()

    if len(parts) < 3:
        return fio

    return f"{parts[0]} {parts[1][0]}.{parts[2][0]}."


def generate_contract(case):

    data = parse_client_data(case)

    now = datetime.now()

    months = {
        1: "января",
        2: "февраля",
        3: "марта",
        4: "апреля",
        5: "мая",
        6: "июня",
        7: "июля",
        8: "августа",
        9: "сентября",
        10: "октября",
        11: "ноября",
        12: "декабря"
    }

    day = now.day
    month = months[now.month]
    year = now.year

    # считаем номер договора

    folder = "cases"
    count = 0

    for root, dirs, files in os.walk(folder):
        for file in files:
            if file.startswith("договор") and str(year)[-2:] in file:
                count += 1

    number = f"{count + 1}/{now.month:02d}-{str(year)[-2:]}"

    signature = make_signature(data["ФИО"])

    replacements = {

        "{{Номер_договора}}": number,
        "{{День}}": str(day),
        "{{Месяц}}": month,
        "{{Год}}": str(year),

        "{{ФИО}}": data["ФИО"],
        "{{Пол}}": data["Пол"],
        "{{Дата_рождения}}": data["Дата рождения"],
        "{{Место_рождения}}": data["Место рождения"],

        "{{Серия}}": data["Серия"],
        "{{Номер}}": data["Номер"],

        "{{Дата_выдачи}}": data["Дата выдачи"],
        "{{Кем_выдан}}": data["Кем выдан"],

        "{{Код_подразделения}}": data["Код подразделения"],

        "{{Адрес_регистрации}}": data["Адрес регистрации"],

        "{{Email}}": data["Email"],
        "{{Телефон}}": data["Телефон"],

        "{{ФИО_подпись}}": signature
    }

    template = os.path.join(os.getcwd(), "templates", "dogovor_template.docx")

    if not os.path.exists(template):
        raise FileNotFoundError("Шаблон договора не найден: templates/dogovor_template.docx")


    doc = Document(template)

    # -----------------------------------
    # замена переменных в обычном тексте
    # -----------------------------------

    for p in doc.paragraphs:

        for run in p.runs:

            for key, value in replacements.items():

                if key in run.text:

                    run.text = run.text.replace(key, value)


    # -----------------------------------
    # замена переменных в таблицах
    # -----------------------------------

    for table in doc.tables:

        for row in table.rows:

            for cell in row.cells:

                for p in cell.paragraphs:

                    for run in p.runs:

                        for key, value in replacements.items():

                            if key in run.text:

                                run.text = run.text.replace(key, value)


    # -----------------------------------
    # сохраняем DOCX
    # -----------------------------------

    case_folder = f"cases/{case}"

    os.makedirs(case_folder, exist_ok=True)

    docx_path = f"{case_folder}/договор {case}.docx"

    doc.save(docx_path)


    # -----------------------------------
    # конвертация DOCX → PDF
    # -----------------------------------

    subprocess.run([
        "/usr/local/bin/soffice",
        "--headless",
        "--convert-to",
        "pdf",
        "--outdir",
        case_folder,
        docx_path
    ], check=True)


    # -----------------------------------
    # путь к PDF
    # -----------------------------------

    pdf_path = f"{case_folder}/договор {case}.pdf"

    return docx_path, pdf_path