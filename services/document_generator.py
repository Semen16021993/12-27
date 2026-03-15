from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_LINE_SPACING
import subprocess
import os


def set_document_margins(doc):
    """Устанавливает поля документа"""
    sections = doc.sections
    for section in sections:
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)


def style_paragraph(p):
    """Применяет стили к абзацу"""
    if not p.runs:
        p.add_run()

    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

    fmt = p.paragraph_format

    # Отступы слева и справа
    fmt.left_indent = Cm(0.8)
    fmt.right_indent = Cm(0)
    
    # Красная строка (отступ первой строки)
    fmt.first_line_indent = Cm(1.25)  # или Cm(0) если не нужна

    # Интервалы
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    
    # Выравнивание
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def replace_placeholders(paragraph, replacements):
    """Заменяет плейсхолдеры в абзаце"""
    for run in paragraph.runs:
        for key, value in replacements.items():
            if key != "{{BODY}}" and key in run.text:
                run.text = run.text.replace(key, value)


def generate_document(case, template_name, replacements):
    """Генерирует документ из шаблона"""
    
    template_path = f"templates/{template_name}"
    doc = Document(template_path)
    
    # Устанавливаем поля документа
    set_document_margins(doc)
    
    body_text = replacements.get("{{BODY}}", "")
    paragraphs = [p.strip() for p in body_text.split("\n\n") if p.strip()]
    
    body_index = None
    
    # Ищем плейсхолдер {{BODY}}
    for i, p in enumerate(doc.paragraphs):
        replace_placeholders(p, replacements)
        if "{{BODY}}" in p.text:
            body_index = i
    
    if body_index is None:
        raise Exception("{{BODY}} not found in template")
    
    # Обрабатываем абзац с плейсхолдером
    body_paragraph = doc.paragraphs[body_index]
    body_paragraph.clear()
    
    if paragraphs:
        body_paragraph.add_run(paragraphs[0])
        style_paragraph(body_paragraph)
    
    # Добавляем остальные абзацы
    insert_index = body_index + 1
    
    for para in paragraphs[1:]:
        if insert_index < len(doc.paragraphs):
            new_p = doc.paragraphs[insert_index].insert_paragraph_before(para)
        else:
            new_p = doc.add_paragraph(para)
        
        style_paragraph(new_p)
        insert_index += 1
    
    # Создаем папку для сохранения
    folder = f"cases/{case}/defense"
    os.makedirs(folder, exist_ok=True)
    
    base_name = template_name.replace(".docx", "")
    docx_path = f"{folder}/{base_name} {case}.docx"
    
    # Сохраняем документ
    doc.save(docx_path)
    
    # Конвертируем в PDF
    try:
        subprocess.run([
            "/usr/local/bin/soffice",
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            folder,
            docx_path
        ], check=True)
    except subprocess.CalledProcessError as e:
        print(f"Ошибка конвертации в PDF: {e}")
    except FileNotFoundError:
        print("LibreOffice не найден. Установите его или проверьте путь.")
    
    pdf_path = f"{folder}/{base_name} {case}.pdf"
    
    return docx_path, pdf_path